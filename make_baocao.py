# -*- coding: utf-8 -*-
"""
Sinh bao cao Word (.docx) DAY DU cho de tai Lap trinh nhan Linux.
- 4 chuong + bia + danh muc hinh + muc luc + danh muc viet tat
- Nhung toan bo ma nguon Phan 1-3 (doc tu thu muc code/) va code GPIO Phan 4
- Cac vi tri can anh chup duoc danh dau bang KHUNG PLACEHOLDER + caption tu dong
"""
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'v:\Project\ltnln'
CODE = os.path.join(ROOT, 'code')

doc = Document()

# ---------- Font mac dinh ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(13)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(6)

# ---------- Tien ich ----------
def _set_shd(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith('}tc') else el.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pr.append(shd)

def body(text, bullet=False, just=True):
    p = doc.add_paragraph(text, style='List Bullet' if bullet else None)
    if not bullet and just:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def sub(text):
    """Dong mo ta phu/ gach dau dong cap 1 (khong phai bullet Word)."""
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def code_block(text):
    """Khoi ma nguon: chu Consolas, nen xam, co the ngat trang."""
    lines = text.rstrip('\n').split('\n')
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        _set_shd(p._p, 'F3F3F3')
        r = p.add_run(ln if ln else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def code_file(relpath):
    with open(os.path.join(CODE, relpath), 'r', encoding='utf-8') as f:
        code_block(f.read())

def _seq_caption(label_text, desc):
    """Tao caption co so thu tu tu dong (SEQ) de Danh muc hinh tu cap nhat."""
    p = doc.add_paragraph(style='Caption') if 'Caption' in [s.name for s in doc.styles] else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(label_text + ' ')
    r.italic = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    # SEQ field
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = ' SEQ Hình \\* ARABIC '
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    tt = OxmlElement('w:t'); tt.text = '1'
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run = p.add_run(); run.italic = True; run.font.size = Pt(12)
    run._r.append(fb); run._r.append(it); run._r.append(fs); run._r.append(tt); run._r.append(fe)
    r2 = p.add_run('. ' + desc)
    r2.italic = True; r2.font.size = Pt(12); r2.font.name = 'Times New Roman'

def figure(desc, hint):
    """Chen 1 KHUNG placeholder anh + caption tu dong.
       desc: ten hinh (vd 'Kết quả chương trình quản lý file')
       hint: huong dan chup gi de dan vao."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    # vien dut net
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'dashed'); e.set(qn('w:sz'), '6'); e.set(qn('w:color'), '999999')
        borders.append(e)
    tbl._tbl.tblPr.append(borders)
    _set_shd(cell._tc, 'FAFAFA')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run('🖼  [ CHÈN ẢNH CHỤP MÀN HÌNH ]')
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Nội dung cần chụp: ' + hint)
    r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p2.paragraph_format.space_after = Pt(18)
    _seq_caption('Hình', desc)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def center(text, size=13, bold=False, italic=False, space=6, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def field(instr, placeholder):
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    tt = OxmlElement('w:t'); tt.text = placeholder
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    for x in (fb, it, fs, tt, fe): run._r.append(x)
    return p

# ============================================================
# TRANG BIA
# ============================================================
center('TRƯỜNG ĐẠI HỌC ............................', 13, True, space=2)
center('KHOA CÔNG NGHỆ THÔNG TIN', 13, True, space=2)
center('--------------------', 13, space=70)
center('BÁO CÁO MÔN HỌC', 16, True, space=4)
center('LẬP TRÌNH NHÂN LINUX', 22, True, space=24, color=RGBColor(0x1F,0x4E,0x79))
center('Đề tài:', 13, italic=True, space=2)
center('LẬP TRÌNH SHELL QUẢN LÝ HỆ THỐNG; LẬP TRÌNH QUẢN LÝ '
       'TIẾN TRÌNH, FILE, SOCKET VÀ NETWORK; XÂY DỰNG MODULE NHÂN '
       'VÀ HỆ THỐNG ARM LINUX MÔ PHỎNG BẰNG BUILDROOT & QEMU', 14, True, space=50)
for line in ['Giảng viên hướng dẫn :  ............................',
             'Sinh viên thực hiện  :  ............................',
             'Mã sinh viên         :  ............................',
             'Lớp                  :  ............................',
             'Khóa                 :  ............................']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run(line).font.size = Pt(13)
center('', space=40)
center('Hà Nội, 2026', 13, True)
doc.add_page_break()

# ============================================================
# DANH MUC HINH ANH
# ============================================================
center('DANH MỤC HÌNH ẢNH', 15, True, space=10)
field(' TOC \\h \\z \\c "Hình" ', 'Nhấn Ctrl+A rồi F9 để cập nhật danh mục hình ảnh (sau khi đã chèn ảnh).')
doc.add_page_break()

# ============================================================
# MUC LUC
# ============================================================
center('MỤC LỤC', 15, True, space=10)
field(' TOC \\o "1-3" \\h \\z \\u ', 'Nhấn Ctrl+A rồi F9 để cập nhật mục lục.')
doc.add_page_break()

# ============================================================
# DANH MUC TU VIET TAT
# ============================================================
center('DANH MỤC KÝ HIỆU VÀ CHỮ VIẾT TẮT', 15, True, space=10)
abbr = [('Từ viết tắt', 'Ý nghĩa'),
        ('OS', 'Operating System – Hệ điều hành'),
        ('LKM', 'Loadable Kernel Module – Module nhân nạp động'),
        ('GPIO', 'General Purpose Input/Output – Cổng vào/ra đa dụng'),
        ('NTP', 'Network Time Protocol – Giao thức đồng bộ thời gian'),
        ('APT', 'Advanced Package Tool – Công cụ quản lý gói của Debian/Ubuntu'),
        ('TCP', 'Transmission Control Protocol'),
        ('IP', 'Internet Protocol'),
        ('QEMU', 'Quick Emulator – Phần mềm giả lập phần cứng'),
        ('ARM', 'Advanced RISC Machine – Kiến trúc vi xử lý'),
        ('GCC', 'GNU Compiler Collection'),
        ('PID', 'Process Identifier – Định danh tiến trình'),
        ('SSH', 'Secure Shell')]
t = doc.add_table(rows=len(abbr), cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b) in enumerate(abbr):
    c0, c1 = t.rows[i].cells
    c0.text = a; c1.text = b
    if i == 0:
        for c in (c0, c1):
            c.paragraphs[0].runs[0].bold = True
            _set_shd(c._tc, 'D9E2F3')
doc.add_page_break()

# ============================================================
# LOI MO DAU
# ============================================================
h1('LỜI MỞ ĐẦU')
body('Trong thời đại bùng nổ công nghệ thông tin hiện nay, các phần mềm mã nguồn '
     'mở ngày càng đóng vai trò quan trọng. Linux là một trong những hệ điều hành '
     'mã nguồn mở tiêu biểu nhất: miễn phí, ổn định, bảo mật cao và có thể tùy biến '
     'sâu. Linux hiện diện ở khắp nơi, từ máy chủ, điện toán đám mây, thiết bị nhúng '
     'cho đến điện thoại di động (nhân Android) và siêu máy tính.')
body('Việc nắm vững cách lập trình tương tác với hệ điều hành Linux — từ lập trình '
     'kịch bản shell, lập trình C sử dụng lời gọi hệ thống (system call), đến lập '
     'trình trong không gian nhân (kernel space) — là kiến thức nền tảng quan trọng '
     'với sinh viên ngành Công nghệ thông tin.')
body('Báo cáo trình bày quá trình tìm hiểu và xây dựng chương trình cho đề tài môn '
     'học “Lập trình nhân Linux”, gồm các nội dung:')
body('Chương 1: Giới thiệu chung về Linux và hệ thống nhúng.', bullet=True)
body('Chương 2: Xây dựng chương trình — lập trình shell quản lý hệ thống và lập '
     'trình C quản lý tiến trình, file, socket, network.', bullet=True)
body('Chương 3: Lập trình xây dựng module nhân và tích hợp vào hệ thống.', bullet=True)
body('Chương 4: Xây dựng hệ thống ARM Linux mô phỏng bằng Buildroot và QEMU.', bullet=True)
body('Em xin chân thành cảm ơn thầy/cô giảng viên đã cung cấp kiến thức và tài liệu '
     'để em hoàn thành báo cáo. Trong quá trình thực hiện không tránh khỏi thiếu sót, '
     'em rất mong nhận được sự góp ý của thầy/cô.')
doc.add_page_break()

# ============================================================
# CHUONG 1
# ============================================================
h1('CHƯƠNG 1. GIỚI THIỆU CHUNG')

h2('1.1. Tổng quan về hệ điều hành Linux')
h3('1.1.1. Khái niệm')
body('Linux là một hệ điều hành mã nguồn mở (open-source), chạy trên hầu hết các '
     'kiến trúc vi xử lý, bao gồm cả dòng vi xử lý ARM. Linux được phát triển và hỗ '
     'trợ bởi cộng đồng mã nguồn mở (GNU). Nhờ tính mở, người dùng được quyền sử '
     'dụng miễn phí, sao chép, nghiên cứu và sửa đổi mã nguồn. Hầu hết các bản Linux '
     'đều hỗ trợ nhiều ngôn ngữ lập trình, đặc biệt công cụ GCC cho phép biên dịch và '
     'thực thi ứng dụng viết bằng C/C++, Java, Python… Ngày nay Linux phân ra nhiều '
     'nhánh (distro) như Ubuntu, Linux Mint, Fedora, CentOS…, trong đó Ubuntu là bản '
     'phân phối thông dụng nhất.')
figure('Hệ điều hành Linux', 'logo/ảnh minh họa hệ điều hành Linux (có thể lấy ảnh minh họa).')

h3('1.1.2. Cấu trúc của hệ điều hành Linux')
body('Hệ điều hành Linux gồm ba thành phần chính:')
sub('► Kernel (Nhân): thành phần quan trọng nhất, chứa các module và thư viện để '
    'quản lý và giao tiếp giữa phần cứng máy tính với các ứng dụng (quản lý tiến '
    'trình, bộ nhớ, thiết bị, hệ thống tệp…).')
sub('► Shell: bộ phận tiếp nhận lệnh từ người dùng hoặc ứng dụng, phiên dịch và '
    'chuyển tới Kernel xử lý. Một số loại shell phổ biến: sh, bash, csh, ash, zsh.')
sub('► Application (Ứng dụng): lớp trên cùng, nơi người dùng cài đặt và chạy các '
    'ứng dụng phục vụ nhu cầu của mình.')
figure('Cấu trúc hệ điều hành Linux', 'sơ đồ phân lớp Kernel – Shell – Application.')

h3('1.1.3. Ưu, nhược điểm của Linux')
body('Ưu điểm:')
sub('► Tính linh hoạt: cho phép chỉnh sửa hệ điều hành theo nhu cầu — lý tưởng cho '
    'lập trình viên và nhà phát triển.')
sub('► Hoạt động tốt trên máy cấu hình thấp, được hỗ trợ cập nhật thường xuyên.')
sub('► Miễn phí bản quyền, đầy đủ tính năng (kèm bộ ứng dụng văn phòng LibreOffice…).')
sub('► Tính bảo mật cao: phần lớn mã độc trên Windows không hoạt động được trên Linux.')
body('Nhược điểm:')
sub('► Số lượng ứng dụng hỗ trợ chưa phong phú bằng Windows.')
sub('► Một số nhà sản xuất không phát triển driver cho Linux.')
sub('► Người dùng cần thời gian làm quen khi chuyển từ Windows sang Linux.')

h3('1.1.4. Vai trò của Linux')
body('Tổ chức, quản lý và vận hành máy chủ, các dịch vụ trong hệ thống thông tin.', bullet=True)
body('Cung cấp công cụ và giải pháp bảo đảm an toàn thông tin.', bullet=True)
body('Là nền tảng chính của điện toán đám mây (Cloud) và dữ liệu lớn (Big Data).', bullet=True)
body('Là nhân cơ bản của hệ điều hành Android trên thiết bị di động.', bullet=True)
body('Được dùng rộng rãi trong các thiết bị nhúng và các giải pháp ảo hóa.', bullet=True)

h2('1.2. Tổng quan về hệ thống nhúng (Embedded Linux)')
h3('1.2.1. Khái niệm')
body('Hệ thống nhúng là tập hợp phần cứng và phần mềm máy tính dựa trên vi điều '
     'khiển hoặc bộ vi xử lý, thường được thiết kế để thực hiện một số chức năng '
     'chuyên biệt. Linux là hệ điều hành được lựa chọn cho nhiều hệ thống nhúng như '
     'điện thoại thông minh, máy tính bảng, set-top box, thiết bị mạng… Linux nhúng '
     'tuy dùng chung nhân nhưng được tinh gọn để có kích thước nhỏ và yêu cầu tài '
     'nguyên thấp hơn so với bản Linux tiêu chuẩn.')
h3('1.2.2. Ưu, nhược điểm')
sub('► Ưu điểm: mã nguồn mở, dễ tùy biến; cấu trúc module, chiếm ít bộ nhớ; độc lập '
    'nền tảng (biên dịch chéo cho nhiều kiến trúc); tái sử dụng được mã nguồn.')
sub('► Nhược điểm: không phải hệ điều hành thời gian thực nên không phù hợp một số '
    'ứng dụng đòi hỏi xử lý khẩn cấp; thiếu một chuẩn thống nhất và không thuộc một '
    'nhà cung cấp duy nhất nên hỗ trợ kỹ thuật còn hạn chế.')
h3('1.2.3. Một số lĩnh vực ứng dụng')
body('Thiết bị nhúng chạy Linux; điện thoại và máy tính bảng; bộ định tuyến mạng; '
     'thiết bị IoT, thiết bị y tế, ô tô, hệ thống định vị, truyền hình – giải trí…')

h2('1.3. Mục tiêu và phạm vi đề tài')
body('Viết các kịch bản shell quản lý file, lập lịch tác vụ, thiết lập thời gian hệ '
     'thống và cài đặt/gỡ bỏ phần mềm tự động.', bullet=True)
body('Viết chương trình C quản lý tiến trình, file, lập trình socket (TCP) và lấy '
     'thông tin mạng bằng các lời gọi hệ thống của Linux.', bullet=True)
body('Xây dựng, biên dịch và nạp một module nhân Linux vào hệ thống.', bullet=True)
body('Xây dựng hệ thống ARM Linux nhúng mô phỏng bằng Buildroot và QEMU.', bullet=True)

h2('1.4. Môi trường và công cụ')
body('Hệ điều hành: Ubuntu 20.04/22.04 (64-bit).', bullet=True)
body('Trình biên dịch GCC, công cụ make; gói build-essential và '
     'linux-headers-$(uname -r) để biên dịch module nhân.', bullet=True)
body('Các tiện ích hệ thống: bash, crontab, apt/dpkg, ntpdate, tar, ps, dmesg.', bullet=True)
body('Buildroot và QEMU cho phần hệ thống nhúng ARM.', bullet=True)
doc.add_page_break()

# ============================================================
# CHUONG 2
# ============================================================
h1('CHƯƠNG 2. XÂY DỰNG CHƯƠNG TRÌNH')
body('Chương này trình bày phần lập trình shell (Phần 1) và lập trình C (Phần 2) để '
     'quản lý hệ thống Linux. Mỗi chương trình được tổ chức dạng menu cho phép người '
     'dùng chọn chức năng bằng số.')

h2('2.1. Lập trình Shell quản lý hệ thống')

h3('2.1.1. Chương trình quản lý file')
body('Chương trình quanly_file.sh cung cấp các chức năng: tạo thư mục/file, đổi tên, '
     'sao chép, di chuyển, xóa, điều hướng, liệt kê, phân quyền (chmod) và nén (tar). '
     'Hàm chính dùng vòng lặp while kết hợp case để điều hướng đến từng chức năng. '
     'Mỗi chức năng nhận tham số nhập từ bàn phím bằng read và thực thi lệnh hệ thống '
     'tương ứng (mkdir, mv, cp, rm, chmod, tar…). Mã nguồn đầy đủ:')
code_file('1_shell/quanly_file.sh')
body('Kết quả thực thi:')
figure('Kết quả thực thi chương trình quản lý file', 'terminal chạy ./quanly_file.sh, thao tác tạo/đổi tên/xóa thư mục và in [OK].')

h3('2.1.2. Chương trình lập lịch tác vụ')
body('Chương trình laplich_tacvu.sh dùng crontab để liệt kê, thêm, xóa từng tác vụ '
     'và xóa toàn bộ tác vụ định kỳ. Cú pháp crontab gồm 5 trường: phút, giờ, ngày, '
     'tháng, thứ, theo sau là lệnh cần chạy.')
code_file('1_shell/laplich_tacvu.sh')
figure('Kết quả thực thi chương trình lập lịch tác vụ', 'terminal thêm một tác vụ crontab rồi liệt kê lại danh sách.')

h3('2.1.3. Chương trình thiết lập thời gian hệ thống')
body('Chương trình thoigian_hethong.sh cho phép xem giờ, đặt giờ/ngày/múi giờ và '
     'đồng bộ thời gian từ máy chủ NTP (ntpdate). Các thao tác thay đổi giờ cần quyền '
     'sudo.')
code_file('1_shell/thoigian_hethong.sh')
figure('Kết quả thực thi chương trình thiết lập thời gian hệ thống', 'terminal xem giờ hiện tại và đồng bộ NTP thành công.')

h3('2.1.4. Chương trình cài đặt và gỡ bỏ phần mềm')
body('Chương trình caidat_goibo.sh tự động kiểm tra, cài đặt, gỡ bỏ, cập nhật và tìm '
     'kiếm gói phần mềm bằng apt và dpkg.')
code_file('1_shell/caidat_goibo.sh')
figure('Kết quả thực thi chương trình cài đặt/gỡ bỏ phần mềm', 'terminal cài một gói (vd: tree) rồi gỡ bỏ gói đó.')

h3('2.1.5. Menu tổng hợp')
body('Kịch bản menu.sh đóng vai trò trình đơn chính, gọi đến bốn chương trình con ở '
     'trên.')
code_file('1_shell/menu.sh')

h2('2.2. Lập trình C quản lý tiến trình, file, socket và network')

h3('2.2.1. Quản lý tiến trình')
body('Chương trình quanly_tientrinh.c minh họa quản lý tiến trình trên Linux: liệt '
     'kê tiến trình (lệnh ps), tạo tiến trình con bằng fork() kết hợp execl(), gửi '
     'tín hiệu và tiêu diệt tiến trình bằng kill(), chờ tiến trình con bằng waitpid().')
code_file('2_c/quanly_tientrinh.c')
figure('Kết quả thực thi chương trình quản lý tiến trình', 'terminal liệt kê tiến trình và tạo một tiến trình con chạy lệnh.')

h3('2.2.2. Quản lý file')
body('Chương trình quanly_file.c dùng các lời gọi hệ thống (stat, chmod, opendir/'
     'readdir) và thư viện chuẩn để kiểm tra tồn tại, tạo, đọc, ghi thêm, xóa, xem/'
     'cấp quyền và liệt kê file trong thư mục.')
code_file('2_c/quanly_file.c')
figure('Kết quả thực thi chương trình quản lý file (C)', 'terminal tạo file, ghi nội dung, đọc lại và xem quyền của file.')

h3('2.2.3. Thông tin mạng (Network)')
body('Chương trình network.c dùng getifaddrs() để liệt kê toàn bộ giao diện mạng '
     'cùng địa chỉ IPv4/IPv6 tương ứng của máy.')
code_file('2_c/network.c')
figure('Kết quả thực thi chương trình network', 'terminal in danh sách giao diện mạng (lo, eth0/ens33…) kèm địa chỉ IP.')

h3('2.2.4. Lập trình Socket (mô hình Client – Server TCP)')
body('Cặp chương trình server.c và client.c minh họa giao tiếp qua socket TCP. Phía '
     'server tạo socket, bind vào cổng 8080, listen và accept; phía client connect '
     'đến server. Hai bên trao đổi tin nhắn, gõ “bye” để kết thúc.')
body('Mã nguồn phía Server:')
code_file('2_c/server.c')
body('Mã nguồn phía Client:')
code_file('2_c/client.c')
body('Toàn bộ chương trình C được biên dịch bằng Makefile sau (lệnh make / make clean):')
code_file('2_c/Makefile')
figure('Kết quả thực thi chương trình Socket', 'hai cửa sổ terminal: server và client nhắn tin qua lại.')
doc.add_page_break()

# ============================================================
# CHUONG 3
# ============================================================
h1('CHƯƠNG 3. LẬP TRÌNH MODULE NHÂN LINUX')
h2('3.1. Khái niệm module nhân')
body('Module nhân (Loadable Kernel Module – LKM) là đoạn mã có thể nạp vào hoặc gỡ '
     'khỏi nhân Linux ngay khi hệ thống đang chạy mà không cần biên dịch lại toàn bộ '
     'nhân. Module thường được dùng để viết driver thiết bị, hệ thống tệp hoặc bổ '
     'sung chức năng cho nhân. Một số lệnh quản lý module: insmod (nạp), rmmod (gỡ), '
     'lsmod (liệt kê), modinfo (xem thông tin).')
h2('3.2. Cấu trúc một module')
body('Mọi module đều có một hàm khởi tạo (đánh dấu __init, đăng ký qua module_init) '
     'chạy khi nạp module, và một hàm kết thúc (đánh dấu __exit, đăng ký qua '
     'module_exit) chạy khi gỡ module. Việc xuất thông tin trong nhân dùng printk '
     '(thay cho printf) và được xem bằng lệnh dmesg. Module có thể nhận tham số đầu '
     'vào qua macro module_param. Giấy phép, tác giả, mô tả được khai báo bằng các '
     'macro MODULE_LICENSE, MODULE_AUTHOR, MODULE_DESCRIPTION.')
h2('3.3. Chương trình minh họa')
body('Module hello_module.c nhận tham số đầu vào (n và ten), tính tổng các số từ 1 '
     'đến n và liệt kê các số nguyên tố ≤ n ngay trong không gian nhân, sau đó ghi '
     'kết quả ra nhật ký nhân bằng printk.')
code_file('3_kernel/hello_module.c')
h2('3.4. Biên dịch và nạp module')
body('Module được biên dịch dựa trên hệ thống Kbuild của nhân thông qua Makefile:')
code_file('3_kernel/Makefile')
body('Các bước biên dịch, nạp và gỡ module:')
code_block('''make                                   # tao hello_module.ko
sudo insmod hello_module.ko n=10 ten="Linux"
dmesg | tail -n 20                     # xem log cua module
sudo rmmod hello_module                # go module
make clean''')
body('Kết quả ghi trong nhật ký nhân (dmesg) khi nạp module với n=10:')
code_block('''hello_module: Xin chao, Linux!
hello_module: Module da duoc nap. Tham so n=10
hello_module: Tong 1..10 = 55
hello_module: Cac so nguyen to <= 10:
2 3 5 7
hello_module: Co tong cong 4 so nguyen to <= 10''')
figure('Kết quả thực thi chương trình module nhân', 'terminal: insmod module rồi chạy dmesg hiển thị các dòng log ở trên; và dòng "Tam biet" khi rmmod.')
doc.add_page_break()

# ============================================================
# CHUONG 4
# ============================================================
h1('CHƯƠNG 4. HỆ THỐNG ARM LINUX MÔ PHỎNG BẰNG BUILDROOT VÀ QEMU')

h2('4.1. Giới thiệu')
body('Hệ thống ARM Linux mô phỏng bằng Buildroot và QEMU là một môi trường phát '
     'triển linh hoạt để phát triển và kiểm thử ứng dụng trên các hệ thống nhúng dựa '
     'trên kiến trúc ARM mà không cần thiết bị vật lý. Trong đề tài này, hệ thống '
     'được xây dựng cho kiến trúc ARM 64-bit (ARMv8-A, CPU Cortex-A53).')

h2('4.2. Buildroot')
body('Buildroot là công cụ tự động hóa quá trình xây dựng hệ thống nhúng. Nó cho '
     'phép tạo ra một hệ thống nhúng tùy chỉnh từ mã nguồn mở, bao gồm nhân Linux, '
     'thư viện và các ứng dụng cần thiết. Người dùng chọn và cấu hình các gói qua '
     'menuconfig hoặc tệp cấu hình; Buildroot sẽ tự động tải về, biên dịch (kể cả '
     'biên dịch chéo) và triển khai vào hệ thống đích.')
figure('Buildroot', 'logo hoặc trang chủ Buildroot.')
body('Một số tính năng nổi bật của Buildroot:')
sub('+ Tính linh hoạt: tùy chỉnh từ nhân Linux đến ứng dụng và thư viện cụ thể.')
sub('+ Tối ưu kích thước: tạo ra hệ thống nhúng nhỏ gọn, hiệu suất cao.')
sub('+ Tự động hóa: tải mã nguồn, biên dịch, triển khai, giảm thao tác thủ công.')
sub('+ Hỗ trợ nhiều nền tảng phần cứng khác nhau.')
body('Cấu trúc của Buildroot gồm một số thành phần: Makefile (chứa sẵn script build '
     'OS với các cấu hình khác nhau), Config.in (chứa các tùy chọn cấu hình), .config '
     '(lưu lại các tùy chọn đã cài đặt). Sau khi build, ta có đầy đủ các thành phần '
     'cơ bản của một hệ điều hành: Kernel, Root filesystem, Bootloader và các thành '
     'phần phụ trợ khác.')
figure('Cấu trúc Buildroot', 'sơ đồ các thành phần đầu ra của Buildroot (kernel, rootfs, bootloader).')

h2('4.3. QEMU')
body('QEMU (Quick Emulator) là phần mềm ảo hóa/giả lập mã nguồn mở có thể giả lập '
     'CPU và phần cứng. QEMU cho phép chạy các hệ điều hành và ứng dụng không tương '
     'thích với nền tảng phần cứng của máy chủ ngay trên máy chủ đó.')
figure('QEMU', 'logo hoặc giao diện QEMU.')
body('Mục tiêu và tính năng của QEMU:')
sub('+ Tính đa năng: hỗ trợ nhiều kiến trúc CPU và nền tảng máy chủ.')
sub('+ Tính tương thích cao: giả lập x86, ARM, PowerPC, SPARC, MIPS…')
sub('+ Tính linh hoạt: chạy trên Linux, Windows, macOS…')
body('Cấu trúc và hoạt động: QEMU gồm hai thành phần chính là máy ảo (emulator) — '
     'giả lập CPU, bộ nhớ, thiết bị để chạy hệ thống khách — và hypervisor — quản lý, '
     'điều khiển máy ảo, chia sẻ tài nguyên và mạng.')
figure('Cấu trúc QEMU', 'sơ đồ emulator + hypervisor của QEMU.')

h2('4.4. Thực thi chương trình')

h3('4.4.1. Xây dựng hình ảnh Linux bằng Buildroot')
body('Hệ thống Linux nhúng ARM 64-bit cần có: hệ thống tệp gốc, hình ảnh nhân và '
     'ứng dụng người dùng; ngoài ra cần một trình biên dịch chéo và một trình mô '
     'phỏng (QEMU) khi thử nghiệm trên nền tảng không phải ARM. Buildroot tự động hóa '
     'toàn bộ quá trình này. Các bước thực hiện:')
body('Bước 1 — Tải Buildroot và tạo cấu hình mặc định cho QEMU ARM 64-bit:')
code_block('''mkdir ~/elec3607 && cd ~/elec3607
git clone git://git.buildroot.net/buildroot
cd buildroot
make qemu_aarch64_virt_defconfig      # cau hinh san cho QEMU ARM 64-bit
make menuconfig                       # tuy chinh them goi''')
figure('Tải xuống và tạo cấu hình Buildroot', 'terminal chạy git clone buildroot và make qemu_aarch64_virt_defconfig.')
body('Bước 2 — Trong màn hình menuconfig, chọn thêm các gói/tùy chọn phục vụ GPIO:')
sub('+ Kernel  >  Linux Kernel Tools  >  gpio')
sub('+ Target packages  >  Libraries  >  Hardware handling  >  libgpiod  (và bật '
    '“install tools”)')
sub('+ Target packages  >  Networking applications  >  openssh (phục vụ kết nối ssh)')
figure('Cấu hình các gói trong menuconfig', 'màn hình menuconfig chọn libgpiod, gpio tools, openssh.')
body('Bước 3 — Cấu hình nhân để bật driver GPIO của phần cứng giả lập, sau đó lưu '
     'lại:')
code_block('''make linux-menuconfig
# Chon: Device Drivers > GPIO Support >
#       Memory mapped GPIO drivers > PrimeCell PL061 GPIO support''')
figure('Cấu hình driver Memory mapped GPIO (PL061)', 'màn hình linux-menuconfig bật PrimeCell PL061 GPIO support.')
body('Bước 4 — Biên dịch toàn bộ hệ thống (lần đầu mất khoảng 30 phút):')
code_block('''time make''')
body('Bước 5 — Khi build hoàn tất, chạy hệ thống bằng QEMU:')
code_block('''cd output/images
./start-qemu.sh
# start-qemu.sh goi: qemu-system-aarch64 -M virt -cpu cortex-a53 -nographic ...
# Dang nhap: tai khoan "root" (khong can mat khau)
# Thoat QEMU: nhan Ctrl-a roi x''')
figure('Hệ thống Linux ARM khởi động trên QEMU', 'màn hình QEMU boot xong và đăng nhập được vào shell root.')

h3('4.4.2. Cài đặt các gói bổ sung')
body('Số gói cài ở bước trước còn ít nên tính hữu dụng còn hạn chế. Trong thư mục '
     'chính của Buildroot, gõ make menuconfig rồi chọn Target packages > Networking '
     'applications > openssh, sau đó gõ make. Lần build này nhanh hơn nhiều vì chỉ '
     'openssh cần tải, biên dịch và đưa vào hình ảnh hệ thống. Trong môi trường QEMU, '
     'gõ "ssh 10.0.2.2" để kết nối tới máy chủ (QEMU dùng cơ chế user-mode networking '
     'với địa chỉ máy chủ là 10.0.2.2).')

h3('4.4.3. Ứng dụng người dùng: Hello QEMU')
body('Tạo thư mục và viết chương trình hello.c:')
code_block('''#include <stdio.h>

int main(void)
{
    printf("Hello QEMU\\n");
    return 0;
}''')
body('Do mục tiêu là máy ARM khách (không phải máy chủ), ta phải biên dịch bằng '
     'trình biên dịch chéo do Buildroot tạo ra:')
code_block('''~/elec3607/buildroot/output/host/bin/aarch64-linux-gcc -o hello hello.c''')
figure('Biên dịch chương trình hello.c bằng cross-compiler', 'terminal biên dịch hello.c bằng aarch64-linux-gcc.')
figure('Kết quả thực thi hello.c trên QEMU', 'trong QEMU chạy ./hello in ra "Hello QEMU".')

h2('4.5. Giao tiếp GPIO')

h3('4.5.1. Kiểm tra GPIO bằng libgpiod')
body('Để giao tiếp và quản lý các đường GPIO trên hệ thống Linux mô phỏng, ta sử '
     'dụng thư viện libgpiod cùng các công cụ dòng lệnh đã được tích hợp qua Buildroot. '
     'Đăng nhập vào hệ thống khách với tài khoản root rồi thực thi lệnh gpiodetect để '
     'liệt kê các chip GPIO:')
code_block('''# gpiodetect
gpiochip0 [9030000.pl061] (8 lines)''')
body('Kết quả cho thấy hệ thống nhận diện một chip GPIO duy nhất là gpiochip0 tương '
     'ứng bộ điều khiển 9030000.pl061 (PrimeCell PL061) quản lý 8 đường GPIO. Thực '
     'thi lệnh gpioinfo để xem chi tiết trạng thái từng chân:')
code_block('''# gpioinfo
gpiochip0 - 8 lines:
\tline   0:      unnamed       unused   input  active-high
\tline   1:      unnamed       unused   input  active-high
\tline   2:      unnamed       unused   input  active-high
\tline   3:      unnamed       unused   input  active-high
\tline   4:      unnamed       unused   input  active-high
\tline   5:      unnamed       unused   input  active-high
\tline   6:      unnamed       unused   input  active-high
\tline   7:      unnamed       unused   input  active-high''')
body('Toàn bộ 8 chân (line 0–7) đang ở trạng thái chưa đặt tên (unnamed), chưa sử '
     'dụng (unused), cấu hình mặc định là ngõ vào (input) và tích cực mức cao '
     '(active-high).')
figure('Kết quả chạy gpiodetect và gpioinfo', 'terminal trong QEMU chạy hai lệnh gpiodetect và gpioinfo.')

h3('4.5.2. Chương trình blink.c và pl061-change.c')
body('Để kiểm tra khả năng điều khiển xuất tín hiệu số và giám sát thay đổi thanh '
     'ghi trên bộ điều khiển PL061, ta xây dựng hai chương trình: blink.c dùng '
     'libgpiod để đổi trạng thái chân ra, và pl061-change.c truy cập trực tiếp vùng '
     'nhớ vật lý để giám sát.')
body('Chương trình blink.c nhấp nháy chân số 3 của gpiochip0 10 lần với chu kỳ trễ '
     '200ms bằng thư viện libgpiod:')
code_block('''/*
**    blink.c - blink gpiochip0 line 3 with a given delay
*/
#include <stdio.h>
#include <stdlib.h>
#include <unistd.h>
#include <gpiod.h>

#define NSPERIOD    200000000ULL    /* output period in ns */

int
main(int argc, char *argv[])
{
    struct gpiod_chip *output_chip;
    struct gpiod_line *output_line;
    struct timespec delay = {0, NSPERIOD};
    struct timespec rem;
    int line_value = 0;

    if (argc == 2)
        delay.tv_nsec = atoll(argv[1]);

    output_chip = gpiod_chip_open_by_number(0);     /* mo /dev/gpiochip0 */
    output_line = gpiod_chip_get_line(output_chip, 3);   /* lam viec voi chan 3 */
    gpiod_line_request_output(output_line, "blink",
                              GPIOD_LINE_ACTIVE_STATE_HIGH);

    for (int i = 0; i < 10; i++) {
        line_value = !line_value;
        gpiod_line_set_value(output_line, line_value);
        nanosleep(&delay, &rem);
    }
    return 0;
}''')
body('Chương trình pl061-change.c giám sát thay đổi của thanh ghi dữ liệu GPIODATA '
     'thông qua ánh xạ bộ nhớ (mmap) từ tệp /dev/mem tại địa chỉ cơ sở 0x09030000:')
code_block('''/*
**    pl061-change.c - display changes to the GPIO_GPIODATA state
*/
#include <stdio.h>
#include <sys/mman.h>
#include <sys/types.h>
#include <sys/stat.h>
#include <fcntl.h>
#include <time.h>
#include <stdint.h>

#define GPIO0_START_ADDR 0x09030000
#define GPIO0_END_ADDR   0x09030fff
#define GPIO0_SIZE (GPIO0_END_ADDR - GPIO0_START_ADDR)
#define GPIO_GPIODATA    0

double
gethrtime()
{
    struct timespec t;
    int64_t ts;
    clock_gettime(CLOCK_MONOTONIC, &t);
    ts = (int64_t)(t.tv_sec) * 1000000000 + (int64_t)(t.tv_nsec);
    return ts / 1.0e9;
}

int
main()
{
    volatile void *gpio_addr;
    volatile unsigned char *gpio_gpiodata_addr;
    unsigned char c, oldc;

    int fd = open("/dev/mem", O_RDWR);
    gpio_addr = mmap(0, GPIO0_SIZE, PROT_READ | PROT_WRITE,
                     MAP_SHARED, fd, GPIO0_START_ADDR);
    gpio_gpiodata_addr = gpio_addr + GPIO_GPIODATA + 0xff;

    oldc = c = *gpio_gpiodata_addr;
    for (;;) {
        c = *gpio_gpiodata_addr;
        if (c != oldc) {
            oldc = c;
            printf("GPIODATA=%x (t=%fs)\\n", c, gethrtime());
        }
    }
}''')
body('Biên dịch chéo cả hai chương trình trên máy chủ (đặt $H là thư mục chứa '
     'Buildroot):')
code_block('''$H/buildroot/output/host/bin/aarch64-buildroot-linux-gnu-gcc -o blink blink.c \\
    -I$H/buildroot/output/staging/usr/include \\
    -L$H/buildroot/output/staging/usr/lib -lgpiod
$H/buildroot/output/host/bin/aarch64-buildroot-linux-gnu-gcc -o pl061-change pl061-change.c''')
body('Sao chép tệp thực thi vào hệ thống khách qua HTTP bằng wget và cấp quyền chạy:')
code_block('''# wget http://10.0.2.2:8080/blink
# wget http://10.0.2.2:8080/pl061-change
# chmod +x blink pl061-change''')
figure('Quá trình chuyển tệp thực thi sang QEMU qua HTTP', 'terminal trong QEMU dùng wget tải blink và pl061-change.')
body('Chạy chương trình giám sát pl061-change dưới nền, sau đó chạy blink:')
code_block('''# ./pl061-change &
# ./blink
GPIODATA=8 (t=46.833710s)
GPIODATA=0 (t=47.033845s)
GPIODATA=8 (t=47.233977s)
GPIODATA=0 (t=47.434119s)
GPIODATA=8 (t=47.634252s)
GPIODATA=0 (t=47.834385s)
GPIODATA=8 (t=48.034518s)
GPIODATA=0 (t=48.234652s)
GPIODATA=8 (t=48.434785s)
GPIODATA=0 (t=48.634919s)''')
body('Khi blink đổi mức logic trên chân số 3, tiến trình pl061-change phát hiện và '
     'ghi nhận thanh ghi GPIODATA luân phiên giữa mức thấp (0) và mức cao (8 — tương '
     'ứng 2^3 của chân số 3) với độ trễ xấp xỉ 0,2 giây đúng thiết kế.')
figure('Kết quả thay đổi GPIODATA ghi nhận từ pl061-change', 'terminal hiển thị các dòng GPIODATA=8/0 luân phiên như trên.')

h3('4.5.3. Chương trình gpio.c — Đồng bộ đa luồng bằng Mutex')
body('Trong thực tế, nhiều tiến trình/luồng truy cập đồng thời một chân GPIO có thể '
     'gây xung đột và lỗi “Device or resource busy”. Để vừa đọc sự kiện vừa ghi giá '
     'trị trên cùng chân GPIO số 3, ta dùng lập trình đa luồng POSIX kết hợp khóa '
     'loại trừ tương hỗ (Mutex). Chương trình gpio.c gồm luồng ghi (gpio_writer) đổi '
     'trạng thái chân 3 và luồng đọc (gpio_reader) phát hiện sự kiện ngắt sườn trên '
     'cùng chân, đồng bộ qua một đối tượng Mutex chung:')
code_block('''/*
**    gpio.c - Blinky under QEMU and libgpiod (reader + writer + mutex)
*/
#include <stdio.h>
#include <stdlib.h>
#include <pthread.h>
#include <assert.h>
#include <gpiod.h>
#include <time.h>

#define CONSUMER "gpiod-lab"
#define NUM_THREADS 2
#define NSPERIOD    200000000ULL

pthread_mutex_t mutexline;     /* chi reader hoac writer duoc truy cap line */

typedef struct _thread_data_t {
    int tid;
    struct gpiod_line *gpioline;
} thread_data_t;

double gethrtime() {
    struct timespec t; int64_t ts;
    clock_gettime(CLOCK_MONOTONIC, &t);
    ts = (int64_t)(t.tv_sec) * 1000000000 + (int64_t)(t.tv_nsec);
    return ts / 1.0e9;
}

void gpio_reader(struct gpiod_line *line) {
    struct timespec timeout = { 0, NSPERIOD / 100ULL };
    struct timespec shortdelay = { 0, NSPERIOD / 100ULL };
    struct timespec rem;
    int r, v, event = 0;
    for (;;) {
        pthread_mutex_lock(&mutexline);
        if ((r = gpiod_line_request_both_edges_events(line, CONSUMER)))
            goto release;
        r = gpiod_line_event_wait(line, &timeout);
        if (r > 0) {
            v = gpiod_line_get_value(line);
            printf("Event %d: Pin=%d at t=%fs\\n", ++event, v, gethrtime());
        }
    release:
        gpiod_line_release(line);
        pthread_mutex_unlock(&mutexline);
        nanosleep(&shortdelay, &rem);
    }
}

void gpio_writer(struct gpiod_line *line) {
    struct timespec delay = { 0, NSPERIOD / 2 };
    struct timespec rem;
    int v = 0, r;
    for (int i = 0; i < 10; i++) {
        pthread_mutex_lock(&mutexline);
        if ((r = gpiod_line_request_output(line, CONSUMER, v)))
            goto release;
        v = !v;
        gpiod_line_set_value(line, v);
    release:
        gpiod_line_release(line);
        pthread_mutex_unlock(&mutexline);
        nanosleep(&delay, &rem);
    }
}

void *thr_func(void *arg) {
    thread_data_t *data = (thread_data_t *)arg;
    switch (data->tid) {
        case 0:  gpio_writer(data->gpioline); break;   /* writer */
        default: gpio_reader(data->gpioline); break;   /* reader */
    }
    pthread_exit(NULL);
}

int main(int argc, char **argv) {
    pthread_t thr[NUM_THREADS];
    thread_data_t thr_data[NUM_THREADS];
    struct gpiod_chip *output_chip;
    struct gpiod_line *output_line;

    output_chip = gpiod_chip_open_by_number(0);
    output_line = gpiod_chip_get_line(output_chip, 3);

    pthread_mutex_init(&mutexline, NULL);
    for (int i = 0; i < NUM_THREADS; ++i) {
        thr_data[i].tid = i;
        thr_data[i].gpioline = output_line;
        assert(pthread_create(&thr[i], NULL, thr_func, &thr_data[i]) == 0);
    }
    pthread_join(thr[0], NULL);
    fflush(stdout);
    return EXIT_SUCCESS;
}''')
body('Biên dịch chéo gpio.c, liên kết thêm thư viện pthread:')
code_block('''$H/buildroot/output/host/bin/aarch64-buildroot-linux-gnu-gcc -o gpio gpio.c \\
    -I$H/buildroot/output/staging/usr/include \\
    -L$H/buildroot/output/staging/usr/lib -lgpiod -lpthread''')
body('Sao chép vào QEMU và thực thi:')
code_block('''# wget http://10.0.2.2:8080/gpio
# chmod +x gpio
# ./gpio
Event 1: Pin=1 at t=19925.878483s
Event 2: Pin=0 at t=19925.978926s
Event 3: Pin=1 at t=19926.077777s
Event 4: Pin=0 at t=19926.179461s
Event 5: Pin=1 at t=19926.277994s
Event 6: Pin=0 at t=19926.381006s
Event 7: Pin=1 at t=19926.479734s
Event 8: Pin=0 at t=19926.582780s
Event 9: Pin=1 at t=19926.681464s
Event 10: Pin=0 at t=19926.784994s''')
body('Kết quả ghi nhận thành công 10 sự kiện biến đổi trạng thái từ luồng đọc, chân '
     'số 3 thay đổi mức logic xen kẽ giữa thấp (Pin=0) và cao (Pin=1) mà không xảy ra '
     'tranh chấp hay báo lỗi bận. Điều này chứng minh khóa Mutex hoạt động hiệu quả '
     'trong việc phân chia tài nguyên GPIO dùng chung giữa hai luồng.')
figure('Kết quả thực thi chương trình đồng bộ đa luồng gpio', 'terminal trong QEMU chạy ./gpio in ra 10 dòng Event như trên.')
doc.add_page_break()

# ============================================================
# KET LUAN
# ============================================================
h1('KẾT LUẬN')
body('Báo cáo đã hoàn thành đầy đủ các mục tiêu đề ra của đề tài: xây dựng bộ kịch '
     'bản shell quản lý hệ thống (quản lý file, lập lịch tác vụ, thiết lập thời gian, '
     'cài đặt/gỡ bỏ phần mềm); các chương trình C quản lý tiến trình, file, socket và '
     'network; một module nhân Linux hoàn chỉnh; và một hệ thống ARM Linux nhúng mô '
     'phỏng bằng Buildroot/QEMU có khả năng giao tiếp GPIO.')
body('Qua quá trình thực hiện, nhóm đã hiểu rõ hơn cơ chế hoạt động của hệ điều hành '
     'Linux, cách lập trình tương tác với nhân ở cả không gian người dùng và không '
     'gian nhân, cũng như quy trình xây dựng một hệ điều hành nhúng từ mã nguồn. '
     'Hướng phát triển tiếp theo có thể mở rộng module nhân thành driver thiết bị '
     'thực tế và bổ sung giao diện cho các chương trình quản lý.')

# ============================================================
# TAI LIEU THAM KHAO
# ============================================================
h1('TÀI LIỆU THAM KHẢO')
for i, ref in enumerate([
    'The Linux Kernel Documentation – https://www.kernel.org/doc/html/latest/',
    'W. R. Stevens, Advanced Programming in the UNIX Environment, Addison-Wesley.',
    'Robert Love, Linux Kernel Development, 3rd Edition, Addison-Wesley.',
    'Buildroot User Manual – https://buildroot.org/downloads/manual/manual.html',
    'QEMU Documentation – https://www.qemu.org/docs/master/',
    'Emulated ARM Linux with Buildroot and QEMU – '
    'https://phwl.org/2021/emulated-ARM-Linux-with-Buildroot-and-QEMU/',
    'libgpiod – https://git.kernel.org/pub/scm/libs/libgpiod/libgpiod.git/',
], 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.space_after = Pt(4)

out = os.path.join(ROOT, 'BaoCao_LapTrinhNhanLinux.docx')
doc.save(out)
print('Da tao:', out)
